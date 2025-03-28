import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_text(cell, text, font_name='Verdana', font_size=10, alignment='center', vertical_alignment='center'):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # Set alignment
    if alignment == 'justify':
        paragraph.alignment = 3  # Justified
    elif alignment == 'center':
        paragraph.alignment = 1  # Centered

    # Set font to Verdana
    rFonts = run._element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), font_name)

    # Set vertical alignment
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), vertical_alignment)
    tcPr.append(vAlign)

    # Set cell borders to continuous line
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '16')  # Border size
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color
        tcBorders.append(border)
    tcPr.append(tcBorders)

def append_to_header(directory, text_to_append):
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            filepath = os.path.join(directory, filename)
            doc = Document(filepath)
            for section in doc.sections:
                header = section.header
                header.paragraphs[0].text += text_to_append
            doc.save(filepath)

def append_to_footer_full(directory, text_to_append):
    for root, dirs, files in os.walk(directory):
        # Skip the "Obsoletos" folder
        if "Obsoletos" in dirs:
            dirs.remove("Obsoletos")
        for filename in files:
            if filename.endswith(".docx"):
                filepath = os.path.join(root, filename)
                print(filepath)
                doc = Document(filepath)
                for section in doc.sections:
                    footer = section.footer
                    if footer.tables:
                        table = footer.tables[0]  # Adjust the index if needed
                        try:
                            cell = table.cell(1, 1)  # Adjust the row and column index if needed
                            set_cell_text(cell, text_to_append)
                        except IndexError:
                            print(f"\033[92mCell (1, 1) does not exist in {filepath}, skipping.\033[0m")
                        try:
                            cell = table.cell(0, 1)
                            set_cell_text(cell, 'Categoria del documento')
                        except IndexError:
                            print(f"\033[92mCell (0, 1) does not exist in {filepath}, skipping.\033[0m")
                    else:
                        footer.paragraphs[0].text += text_to_append
                doc.save(filepath)

def append_to_footer_once(filepath, text_to_append):
    doc = Document(filepath)
    for section in doc.sections:
        footer = section.footer
        if footer.tables:
            table = footer.tables[0]  # Adjust the index if needed
            try:
                cell = table.cell(1, 1)  # Adjust the row and column index if needed
                set_cell_text(cell, text_to_append)
            except IndexError:
                print(f"\033[92mCell (1, 1) does not exist in {filepath}, skipping.\033[0m")
            try:
                cell = table.cell(0, 1)
                set_cell_text(cell, 'Categoria del documento')
            except IndexError:
                print(f"\033[92mCell (0, 1) does not exist in {filepath}, skipping.\033[0m")
        else:
            footer.paragraphs[0].text += text_to_append
    doc.save(filepath)

directory_path = r'##'  # Path to the folder containing the .docx files (## must be replaced with the actual path)
#file_path = r'##'
text_to_append = ' USO INTERNO '
#append_to_header(directory_path, text_to_append)
append_to_footer_full(directory_path, text_to_append)
#append_to_footer_once(file_path, text_to_append)
